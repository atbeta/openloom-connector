"""Tests for the Runner — poll cycle, push to OpenLoom, result write-back."""

from __future__ import annotations

import io
import json
import socket
import threading
from contextlib import contextmanager
from typing import Any

import httpx
import pytest
import respx

from openloom_connector import Connector, FileEntry
from openloom_connector.config import (
    ConnectorConfig,
    OutboundWebhookConfig,
    WebhookSource,
)
from openloom_connector.runner import Runner, _parse_docx, _parse_spec


class InMemoryConnector(Connector):
    """Connector backed by an in-memory dict — for tests."""

    def __init__(self) -> None:
        self.files: dict[str, bytes] = {}

    def list_inbox(self) -> list[FileEntry]:
        prefix = self.inbox_dir.rstrip("/") + "/"
        return [
            FileEntry(path=p, name=p.rsplit("/", 1)[-1], size=len(b))
            for p, b in self.files.items() if p.startswith(prefix)
        ]

    def download(self, path: str) -> bytes | None:
        return self.files.get(path)

    def upload(self, path: str, content: bytes) -> None:
        self.files[path] = content

    def delete_inbox(self, path: str) -> None:
        self.files.pop(path, None)


def _config(connector_class: type[Connector] = InMemoryConnector, **kwargs: Any) -> ConnectorConfig:
    poll = kwargs.pop("poll_interval_seconds", 1)
    inbox = kwargs.pop("inbox_dir", "/inbox")
    outbox = kwargs.pop("outbox_dir", "/outbox")
    archive = kwargs.pop("archive_dir", "")
    secret = kwargs.pop("signing_secret", "")
    prefix = kwargs.pop("task_prefix", "task-")
    return ConnectorConfig(
        connector_class=connector_class,
        connector_kwargs=kwargs,
        openloom_url="http://loom:55413",
        webhook=WebhookSource(url="http://loom:55413/api/webhooks/generic", signing_secret=secret),
        inbox_dir=inbox,
        outbox_dir=outbox,
        archive_dir=archive,
        poll_interval_seconds=poll,
        task_prefix=prefix,
    )


def _build_docx(rows: list[tuple[str, str]]) -> bytes:
    """Build a minimal docx with a single table containing *rows*."""
    from docx import Document

    doc = Document()
    table = doc.add_table(rows=0, cols=2)
    for k, v in rows:
        row = table.add_row().cells
        row[0].text = k
        row[1].text = v
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── _parse_spec ──────────────────────────────────────────────────────────


def test_parse_json() -> None:
    assert _parse_spec(b'{"goal":"x"}', "task-t.json") == {"goal": "x"}


def test_parse_yaml() -> None:
    assert _parse_spec(b"goal: y\n", "task-t.yaml") == {"goal": "y"}


def test_parse_invalid_returns_none() -> None:
    assert _parse_spec(b"not json", "task-t.json") is None
    assert _parse_spec(b":invalid: yaml:", "task-t.yaml") is None
    assert _parse_spec(b"not a dict", "task-t.json") is None


def test_parse_docx_basic() -> None:
    raw = _build_docx([
        ("goal", "fix login CSS"),
        ("workspace", "/Users/me/proj"),
        ("name", "CSS Fix"),
    ])
    spec = _parse_spec(raw, "task-t.docx")
    assert spec == {
        "goal": "fix login CSS",
        "workspace": "/Users/me/proj",
        "name": "CSS Fix",
    }


def test_parse_docx_session_id_and_metadata() -> None:
    raw = _build_docx([
        ("goal", "continue review"),
        ("sessionId", "ses_abc"),
        ("branch", "main"),
        ("priority", "high"),
    ])
    spec = _parse_spec(raw, "task-t.docx")
    assert spec is not None
    assert spec["goal"] == "continue review"
    assert spec["sessionId"] == "ses_abc"
    assert spec["metadata"] == {"branch": "main", "priority": "high"}


def test_parse_docx_empty_returns_none() -> None:
    raw = _build_docx([])
    assert _parse_docx(raw) is None


def test_parse_docx_ignores_empty_rows() -> None:
    raw = _build_docx([("", ""), ("goal", "x"), ("", "y")])
    spec = _parse_docx(raw)
    assert spec == {"goal": "x"}


# ── prefix filtering ──────────────────────────────────────────────────────


def test_poll_skips_files_without_prefix() -> None:
    """Non-task files (no `task-` prefix) are ignored — the inbox may
    contain arbitrary docs uploaded by teammates."""
    conn = InMemoryConnector()
    conn.files["/inbox/random-doc.json"] = json.dumps({"goal": "x"}).encode()
    runner = Runner(_config())
    runner._connector = conn

    with respx.mock:
        route = respx.post("http://loom:55413/api/webhooks/generic").mock(
            return_value=httpx.Response(200, json={"taskId": "t1"}),
        )
        runner._poll_once()

    assert route.call_count == 0


def test_poll_custom_prefix() -> None:
    conn = InMemoryConnector()
    conn.files["/inbox/loom-fix.json"] = json.dumps({"goal": "x", "workspace": "/p"}).encode()
    runner = Runner(_config(task_prefix="loom-"))
    runner._connector = conn

    with respx.mock:
        route = respx.post("http://loom:55413/api/webhooks/generic").mock(
            return_value=httpx.Response(200, json={"taskId": "t1"}),
        )
        runner._poll_once()

    assert route.call_count == 1


# ── poll cycle ───────────────────────────────────────────────────────────


def test_poll_dispatches_new_files() -> None:
    conn = InMemoryConnector()
    conn.files["/inbox/task-task1.json"] = json.dumps({
        "goal": "fix bug", "workspace": "/p",
    }).encode()
    runner = Runner(_config())
    runner._connector = conn

    with respx.mock:
        respx.post("http://loom:55413/api/webhooks/generic").mock(
            return_value=httpx.Response(200, json={"ok": True, "taskId": "task_abc"}),
        )
        runner._poll_once()

    assert "task_abc" in runner._task_to_file
    assert "/inbox/task-task1.json" in runner._seen


def test_poll_skips_seen_files() -> None:
    conn = InMemoryConnector()
    conn.files["/inbox/task-t.json"] = json.dumps({"goal": "x", "workspace": "/p"}).encode()
    runner = Runner(_config())
    runner._connector = conn

    with respx.mock:
        route = respx.post("http://loom:55413/api/webhooks/generic").mock(
            return_value=httpx.Response(200, json={"taskId": "t1"}),
        )
        runner._poll_once()
        runner._poll_once()

    assert route.call_count == 1


def test_poll_skips_files_without_goal() -> None:
    conn = InMemoryConnector()
    conn.files["/inbox/task-bad.json"] = json.dumps({"name": "no goal", "workspace": "/p"}).encode()
    runner = Runner(_config())
    runner._connector = conn

    with respx.mock:
        route = respx.post("http://loom:55413/api/webhooks/generic").mock(
            return_value=httpx.Response(200, json={"taskId": "t1"}),
        )
        runner._poll_once()

    assert route.call_count == 0
    assert "/inbox/task-bad.json" in runner._seen


def test_poll_skips_files_without_workspace_or_session() -> None:
    conn = InMemoryConnector()
    conn.files["/inbox/task-t.json"] = json.dumps({"goal": "no workspace"}).encode()
    runner = Runner(_config())
    runner._connector = conn

    with respx.mock:
        route = respx.post("http://loom:55413/api/webhooks/generic").mock(
            return_value=httpx.Response(200, json={"taskId": "t1"}),
        )
        runner._poll_once()

    assert route.call_count == 0


def test_poll_accepts_session_id_without_workspace() -> None:
    conn = InMemoryConnector()
    conn.files["/inbox/task-t.json"] = json.dumps({
        "goal": "continue", "sessionId": "ses_xyz",
    }).encode()
    runner = Runner(_config())
    runner._connector = conn

    with respx.mock:
        route = respx.post("http://loom:55413/api/webhooks/generic").mock(
            return_value=httpx.Response(200, json={"taskId": "t1"}),
        )
        runner._poll_once()

    assert route.call_count == 1
    body = json.loads(route.calls.last.request.content)
    assert body["sessionId"] == "ses_xyz"


def test_poll_skips_non_task_extensions() -> None:
    conn = InMemoryConnector()
    conn.files["/inbox/task-readme.txt"] = b"text file"
    runner = Runner(_config())
    runner._connector = conn

    with respx.mock:
        route = respx.post("http://loom:55413/api/webhooks/generic").mock(
            return_value=httpx.Response(200, json={"taskId": "t1"}),
        )
        runner._poll_once()

    assert route.call_count == 0


def test_poll_dispatches_docx() -> None:
    conn = InMemoryConnector()
    raw = _build_docx([("goal", "from docx"), ("workspace", "/p")])
    conn.files["/inbox/task-from-docx.docx"] = raw
    runner = Runner(_config())
    runner._connector = conn

    with respx.mock:
        route = respx.post("http://loom:55413/api/webhooks/generic").mock(
            return_value=httpx.Response(200, json={"taskId": "task_docx"}),
        )
        runner._poll_once()

    assert route.call_count == 1
    body = json.loads(route.calls.last.request.content)
    assert body["goal"] == "from docx"


# ── write_result (json input) ────────────────────────────────────────────


def test_write_result_uploads_and_deletes() -> None:
    conn = InMemoryConnector()
    conn.files["/inbox/task-t.json"] = json.dumps({"goal": "x", "workspace": "/p"}).encode()
    runner = Runner(_config())
    runner._connector = conn
    runner._task_to_file["task_xyz"] = "/inbox/task-t.json"
    runner._seen.add("/inbox/task-t.json")

    runner.write_result("task_xyz", "t", "completed", {"summary": "ok"})

    out_files = [p for p in conn.files if p.startswith("/outbox/")]
    assert out_files == ["/outbox/task-t.result.json"]
    result = json.loads(conn.files["/outbox/task-t.result.json"])
    assert result["task_id"] == "task_xyz"
    assert result["status"] == "completed"
    assert result["data"] == {"summary": "ok"}
    assert "/inbox/task-t.json" not in conn.files


def test_write_result_with_archive() -> None:
    conn = InMemoryConnector()
    conn.files["/inbox/task-t.json"] = json.dumps({"goal": "x", "workspace": "/p"}).encode()
    runner = Runner(_config(archive_dir="/archive"))
    runner._connector = conn
    runner._task_to_file["task_xyz"] = "/inbox/task-t.json"

    runner.write_result("task_xyz", "t", "completed", {})

    assert "/inbox/task-t.json" not in conn.files  # deleted
    archive_files = [p for p in conn.files if p.startswith("/archive/")]
    assert archive_files == ["/archive/task-t.json"]


def test_write_result_unknown_task_noop() -> None:
    conn = InMemoryConnector()
    runner = Runner(_config())
    runner._connector = conn

    runner.write_result("nonexistent_task", "t", "completed", {})

    assert conn.files == {}


# ── write_result (docx input → docx output) ─────────────────────────────


def test_write_result_docx_round_trip() -> None:
    conn = InMemoryConnector()
    raw = _build_docx([("goal", "from docx"), ("workspace", "/p")])
    conn.files["/inbox/task-doc.docx"] = raw
    runner = Runner(_config(archive_dir="/archive"))
    runner._connector = conn
    runner._task_to_file["task_doc"] = "/inbox/task-doc.docx"

    runner.write_result("task_doc", "doc task", "completed", {"summary": "done"})

    # Result should be a docx
    out_files = [p for p in conn.files if p.startswith("/outbox/")]
    assert out_files == ["/outbox/task-doc.result.docx"]
    result_bytes = conn.files["/outbox/task-doc.result.docx"]
    assert result_bytes[:2] == b"PK"  # docx is a zip

    # Parse it back — new layout: heading + metadata block + summary section
    from docx import Document
    doc = Document(io.BytesIO(result_bytes))
    headings = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]
    assert "任务完成报告" in headings
    assert "doc task" in headings
    assert "概要" in headings
    body_texts = [p.text for p in doc.paragraphs]
    assert any("done" in t for t in body_texts)

    # Archive + delete
    assert "/inbox/task-doc.docx" not in conn.files
    assert "/archive/task-doc.docx" in conn.files


def test_docx_result_renders_recent_activity_as_separate_blocks() -> None:
    """recent_activity entries must render as distinct sections, and the
    tools inside each entry must be visually separated from the agent
    prose — the JSON-dump blob we used to write was unreadable."""
    from openloom_connector.runner import _render_docx_result

    result = {
        "schema_version": "1.0",
        "task_id": "task_abc",
        "task_name": "登录 CSS 修复",
        "status": "completed",
        "timestamp": 1_700_000_000.0,
        "data": {
            "summary": "Agent reported TASK COMPLETE",
            "recent_activity": [
                {
                    "text": "我看了一下登录页面的样式。",
                    "completed_at": 1_700_000_010.0,
                    "tools": [
                        {
                            "tool": "bash",
                            "status": "completed",
                            "input_excerpt": "ls -la",
                        },
                    ],
                },
                {
                    "text": "改好了。",
                    "completed_at": 1_700_000_020.0,
                    "tools": [
                        {
                            "tool": "edit",
                            "status": "completed",
                            "input_excerpt": "login.css",
                        },
                        {
                            "tool": "bash",
                            "status": "completed",
                            "input_excerpt": "pytest",
                        },
                    ],
                },
            ],
            "active_session_id": "ses_xyz",
        },
    }

    from docx import Document
    doc = Document(_render_docx_result(result))
    all_text = "\n".join(p.text for p in doc.paragraphs)

    # Summary, agent text, and tool excerpts all present and on their own lines
    assert "Agent reported TASK COMPLETE" in all_text
    assert "我看了一下登录页面的样式" in all_text
    assert "改好了" in all_text
    assert "bash" in all_text
    assert "edit" in all_text
    assert "login.css" in all_text

    # "Agent 执行轨迹" section heading present
    headings = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]
    assert any("Agent 执行轨迹" in h for h in headings)

    # Tool lines are separate paragraphs from the agent text (not concatenated)
    agent_paras = [p.text for p in doc.paragraphs if p.text and not p.style.name.startswith("Heading")]
    tool_paras = [t for t in agent_paras if t.lstrip().startswith(("✓", "▶", "…", "✗", "•"))]
    assert len(tool_paras) >= 3, f"expected tool lines as separate paragraphs, got {tool_paras}"


def test_docx_result_handles_failed_status_heading() -> None:
    from openloom_connector.runner import _render_docx_result

    doc = _render_docx_result({
        "status": "failed",
        "task_name": "broken task",
        "task_id": "task_x",
        "timestamp": 0,
        "data": {},
    })
    from docx import Document
    d = Document(doc)
    headings = [p.text for p in d.paragraphs if p.style.name.startswith("Heading")]
    assert "任务失败报告" in headings


def test_docx_result_no_recent_activity_omits_trace_section() -> None:
    from openloom_connector.runner import _render_docx_result

    doc = _render_docx_result({
        "status": "completed",
        "task_name": "minimal task",
        "task_id": "task_y",
        "timestamp": 1_700_000_000.0,
        "data": {"summary": "ok"},
    })
    from docx import Document
    d = Document(doc)
    headings = [p.text for p in d.paragraphs if p.style.name.startswith("Heading")]
    assert not any("Agent 执行轨迹" in h for h in headings)


def test_docx_result_does_not_truncate_long_metadata() -> None:
    """Long tool input_excerpt and other metadata must be preserved in
    full — the agent's execution trace is meaningful content, not a
    summary. Word/WPS handle long-line wrap on the phone."""
    long_input = "a" * 2000
    from openloom_connector.runner import _render_docx_result

    result = {
        "status": "completed",
        "task_name": "long task",
        "task_id": "task_long",
        "timestamp": 1_700_000_000.0,
        "data": {
            "summary": "done",
            "recent_activity": [
                {
                    "text": "ran a big command",
                    "completed_at": 1_700_000_010.0,
                    "tools": [
                        {
                            "tool": "bash",
                            "status": "completed",
                            "input_excerpt": long_input,
                        },
                    ],
                },
            ],
            "active_session_id": long_input,
        },
    }
    from docx import Document
    doc = Document(_render_docx_result(result))
    all_text = "\n".join(p.text for p in doc.paragraphs)
    assert long_input in all_text, "tool input_excerpt should be preserved verbatim"
    # Also verify the leftover metadata block carries the full session id
    # (it lives in a table, not a paragraph).
    table_text = "\n".join(
        cell.text
        for table in doc.tables
        for row in table.rows
        for cell in row.cells
    )
    assert long_input in table_text, "active_session_id should be preserved verbatim"


# ── webhook signing ─────────────────────────────────────────────────────


def test_webhook_signs_when_secret_set() -> None:
    conn = InMemoryConnector()
    conn.files["/inbox/task-t.json"] = json.dumps({"goal": "x", "workspace": "/p"}).encode()
    runner = Runner(_config(signing_secret="topsecret"))
    runner._connector = conn

    with respx.mock:
        route = respx.post("http://loom:55413/api/webhooks/generic").mock(
            return_value=httpx.Response(200, json={"taskId": "t1"}),
        )
        runner._poll_once()

    sig = route.calls.last.request.headers.get("X-OpenLoom-Signature-256")
    assert sig is not None
    assert sig.startswith("sha256=")
    assert len(sig) == len("sha256=") + 64


def test_webhook_no_signature_when_no_secret() -> None:
    conn = InMemoryConnector()
    conn.files["/inbox/task-t.json"] = json.dumps({"goal": "x", "workspace": "/p"}).encode()
    runner = Runner(_config())
    runner._connector = conn

    with respx.mock:
        route = respx.post("http://loom:55413/api/webhooks/generic").mock(
            return_value=httpx.Response(200, json={"taskId": "t1"}),
        )
        runner._poll_once()

    assert "X-OpenLoom-Signature-256" not in route.calls.last.request.headers


# ── run loop ─────────────────────────────────────────────────────────────


async def test_run_loop_stops_on_signal() -> None:
    conn = InMemoryConnector()
    runner = Runner(_config(poll_interval_seconds=0))
    runner._connector = conn

    async def stop_soon() -> None:
        import asyncio
        await asyncio.sleep(0.05)
        runner.stop()

    import asyncio
    await asyncio.gather(runner.run(), stop_soon())


# ── trust_env=False (httpx proxy hardening) ──────────────────────────────────


def test_push_ignores_system_proxy(monkeypatch: pytest.MonkeyPatch) -> None:
    """If HTTP_PROXY is set, the connector must still POST directly to
    127.0.0.1. Regression test for the 'Content Filter - Access Denied'
    symptom caused by httpx honouring system proxy env vars."""
    conn = InMemoryConnector()
    conn.files["/inbox/task-t.json"] = json.dumps({"goal": "x", "workspace": "/p"}).encode()
    runner = Runner(_config())
    runner._connector = conn

    monkeypatch.setenv("HTTP_PROXY", "http://proxy.invalid:9999")
    monkeypatch.setenv("HTTPS_PROXY", "http://proxy.invalid:9999")
    monkeypatch.setenv("ALL_PROXY", "http://proxy.invalid:9999")

    # respx mocks the loopback address, not the proxy. If trust_env is on,
    # httpx routes through the proxy and respx sees nothing. If trust_env
    # is off, the mock intercepts the request directly.
    with respx.mock:
        route = respx.post("http://loom:55413/api/webhooks/generic").mock(
            return_value=httpx.Response(200, json={"taskId": "t1"}),
        )
        runner._poll_once()

    assert route.call_count == 1


# ── outbound webhook receiver ───────────────────────────────────────────────


def _free_port() -> int:
    with socket.socket() as s:
        s.bind(("127.0.0.1", 0))
        return s.getsockname()[1]


@contextmanager
def _running_server(runner: Runner):
    """Start a receiver server in a background thread and yield its port.

    The thread is joined and the listening socket is closed on exit. Use
    this in every receiver test that issues real HTTP requests — building
    the server alone does not call ``accept()``.
    """
    server = runner._build_receiver_server()
    runner._receiver_server = server  # mirrors what ``run()`` sets
    thread = threading.Thread(target=server.serve_forever, daemon=True)
    thread.start()
    try:
        yield server.server_address[1]
    finally:
        server.shutdown()
        server.server_close()
        thread.join(timeout=5)
        runner._receiver_server = None


def _receiver_config(**kwargs: Any) -> ConnectorConfig:
    port = _free_port()
    return ConnectorConfig(
        connector_class=InMemoryConnector,
        connector_kwargs={},
        openloom_url="http://loom:55413",
        webhook=WebhookSource(url="http://loom:55413/api/webhooks/generic"),
        inbox_dir=kwargs.get("inbox_dir", "/inbox"),
        outbox_dir=kwargs.get("outbox_dir", "/outbox"),
        archive_dir=kwargs.get("archive_dir", ""),
        poll_interval_seconds=0,
        task_prefix=kwargs.get("task_prefix", "task-"),
        outbound=OutboundWebhookConfig(
            enabled=True,
            host="127.0.0.1",
            port=port,
            path="/listener/openloom",
        ),
    )


def test_receiver_writes_result_on_task_completed() -> None:
    conn = InMemoryConnector()
    conn.files["/inbox/task-t.json"] = json.dumps({
        "goal": "fix bug", "workspace": "/p",
    }).encode()
    runner = Runner(_receiver_config())
    runner._connector = conn

    with _running_server(runner) as port:
        runner._task_to_file["task_xyz"] = "/inbox/task-t.json"
        runner._seen.add("/inbox/task-t.json")

        payload = {
            "event": "TASK_COMPLETED",
            "task_id": "task_xyz",
            "task_name": "fix bug",
            "data": {"status": "completed", "summary": "done"},
        }
        with httpx.Client(timeout=5.0) as client:
            r = client.post(
                f"http://127.0.0.1:{port}/listener/openloom",
                json=payload,
            )
        assert r.status_code == 200, r.text
        assert r.json()["ok"] is True

        out_files = [p for p in conn.files if p.startswith("/outbox/")]
        assert out_files == ["/outbox/task-t.result.json"]
        result = json.loads(conn.files["/outbox/task-t.result.json"])
        assert result["task_id"] == "task_xyz"
        assert result["status"] == "completed"
        assert "/inbox/task-t.json" not in conn.files


def test_receiver_ignores_unknown_task() -> None:
    conn = InMemoryConnector()
    runner = Runner(_receiver_config())
    runner._connector = conn

    with _running_server(runner) as port:
        with httpx.Client(timeout=5.0) as client:
            r = client.post(
                f"http://127.0.0.1:{port}/listener/openloom",
                json={"event": "TASK_COMPLETED", "task_id": "ghost", "data": {}},
            )
        assert r.status_code == 200
        assert r.json()["ignored"] == "unknown task"
        assert conn.files == {}


def test_receiver_ignores_non_terminal_events() -> None:
    conn = InMemoryConnector()
    conn.files["/inbox/task-t.json"] = json.dumps({"goal": "x", "workspace": "/p"}).encode()
    runner = Runner(_receiver_config())
    runner._connector = conn
    runner._task_to_file["task_xyz"] = "/inbox/task-t.json"

    with _running_server(runner) as port:
        with httpx.Client(timeout=5.0) as client:
            for event_name in ("TASK_CREATED", "TASK_STARTED", "TASK_UPDATED"):
                r = client.post(
                    f"http://127.0.0.1:{port}/listener/openloom",
                    json={
                        "event": event_name,
                        "task_id": "task_xyz",
                        "data": {"status": "running"},
                    },
                )
                assert r.status_code == 200
                assert r.json()["ignored"] == event_name
        # No result file written yet
        assert not any(p.startswith("/outbox/") for p in conn.files)
        # Input still in place (write_result only fires on terminal events)
        assert "/inbox/task-t.json" in conn.files


def test_receiver_404_on_wrong_path() -> None:
    conn = InMemoryConnector()
    runner = Runner(_receiver_config())
    runner._connector = conn

    with _running_server(runner) as port:
        with httpx.Client(timeout=5.0) as client:
            r = client.post(
                f"http://127.0.0.1:{port}/wrong-path",
                json={"event": "TASK_COMPLETED", "task_id": "x"},
            )
        assert r.status_code == 404


def test_receiver_400_on_invalid_body() -> None:
    conn = InMemoryConnector()
    runner = Runner(_receiver_config())
    runner._connector = conn

    with _running_server(runner) as port:
        with httpx.Client(timeout=5.0) as client:
            r = client.post(
                f"http://127.0.0.1:{port}/listener/openloom",
                content=b"not json",
            )
        assert r.status_code == 400


async def test_run_with_receiver_starts_and_stops() -> None:
    """Integration smoke: run() with receiver enabled must accept a real
    inbound HTTP request, and stop() must shut down the receiver thread."""
    import asyncio

    conn = InMemoryConnector()
    conn.files["/inbox/task-t.json"] = json.dumps({
        "goal": "x", "workspace": "/p",
    }).encode()
    runner = Runner(_receiver_config())
    runner._connector = conn
    port = runner._config.outbound.port

    async def driver() -> None:
        # Wait for the server to bind, then poke it.
        for _ in range(50):
            try:
                with httpx.Client(timeout=1.0) as c:
                    r = c.post(
                        f"http://127.0.0.1:{port}/listener/openloom",
                        json={
                            "event": "TASK_COMPLETED",
                            "task_id": "will-be-ignored",
                            "data": {"status": "completed"},
                        },
                    )
                assert r.status_code == 200
                return
            except Exception:
                await asyncio.sleep(0.05)
        raise AssertionError("receiver never came up")

    runner_task = asyncio.create_task(runner.run())
    try:
        await driver()
    finally:
        runner.stop()
        await asyncio.wait_for(runner_task, timeout=5.0)
    assert runner._receiver_server is None
