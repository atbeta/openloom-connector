"""
Runner — polls a Connector, pushes tasks to OpenLoom, writes back results.

Does not import from OpenLoom internals. Talks to OpenLoom only through:
- HTTP POST to ``/api/webhooks/{source}`` (push task)
- (Future) SSE on ``/api/events`` (listen for task completion)

Task file conventions:
- File name must start with ``task_prefix`` (default ``task-``); non-task
  files (e.g. random docs uploaded to the inbox) are ignored.
- Supported formats: ``.json``, ``.yaml``, ``.yml``, ``.docx``.
- For ``.docx``, the spec is read from a 2-column table (field | value).
- Results are written in the same format as the input (``task-X.json`` →
  ``task-X.result.json``, ``task-X.docx`` → ``task-X.result.docx``).
"""

from __future__ import annotations

import asyncio
import hashlib
import hmac
import io
import json
import logging
import time
from http.server import BaseHTTPRequestHandler, ThreadingHTTPServer
from pathlib import PurePosixPath
from typing import Any

import httpx

from .base import Connector, FileEntry
from .config import ConnectorConfig

_logger = logging.getLogger("openloom_connector.runner")

# File extensions that can carry a task spec.
_TASK_EXTENSIONS = {".json", ".yaml", ".yml", ".docx"}

# Map file extension → result suffix (extension).
_RESULT_SUFFIX_BY_EXT = {
    ".json": ".result.json",
    ".yaml": ".result.yaml",
    ".yml": ".result.yml",
    ".docx": ".result.docx",
}


class Runner:
    """Poll-and-forward runner — orchestrates a Connector with OpenLoom."""

    def __init__(self, config: ConnectorConfig) -> None:
        self._config = config
        self._connector: Connector = config.connector_class(**config.connector_kwargs)
        self._seen: set[str] = set()
        self._task_to_file: dict[str, str] = {}
        self._stopped = asyncio.Event()
        # Held so the receiver coroutine can call ``shutdown()`` from another
        # thread when the runner is asked to stop. ``None`` while the receiver
        # is not running.
        self._receiver_server: ThreadingHTTPServer | None = None

    def stop(self) -> None:
        self._stopped.set()
        # Tell the blocking serve_forever thread to return. shutdown() is
        # safe to call from any thread and must be invoked after stop() for
        # ``run()`` to exit cleanly when the receiver is enabled.
        server = self._receiver_server
        if server is not None:
            server.shutdown()

    async def run(self) -> None:
        """Run the polling loop (and outbound receiver if enabled) until
        ``stop()`` is called."""
        _logger.info(
            "connector started — polling every %ds, prefix=%r",
            self._config.poll_interval_seconds,
            self._config.task_prefix,
        )
        _logger.info("openloom url: %s", self._config.openloom_url)
        _logger.info("webhook:      %s", self._config.webhook.url)
        _logger.info("inbox:        %s", self._config.inbox_dir)
        _logger.info("archive:      %s", self._config.archive_dir or "(disabled)")

        tasks: list[asyncio.Task[Any]] = [
            asyncio.create_task(self._poll_loop(), name="openloom-connector-poll"),
        ]
        if self._config.outbound.enabled:
            _logger.info(
                "receiver:     http://%s:%d%s (OpenLoom -> connector)",
                self._config.outbound.host,
                self._config.outbound.port,
                self._config.outbound.path,
            )
            tasks.append(
                asyncio.create_task(
                    self._receiver_loop(), name="openloom-connector-receiver",
                )
            )
        else:
            _logger.info(
                "receiver:     disabled (set outbound_webhook.enabled: true to "
                "listen for task completion events)",
            )

        try:
            await asyncio.gather(*tasks)
        finally:
            for t in tasks:
                if not t.done():
                    t.cancel()
            for t in tasks:
                try:
                    await t
                except (asyncio.CancelledError, Exception):
                    pass

    async def _poll_loop(self) -> None:
        """Inbox polling loop, separated so ``run()`` can ``gather`` it."""
        while not self._stopped.is_set():
            try:
                self._poll_once()
            except Exception:
                _logger.exception("poll cycle failed")
            try:
                await asyncio.wait_for(
                    self._stopped.wait(),
                    timeout=self._config.poll_interval_seconds,
                )
            except TimeoutError:
                pass

    # --- outbound webhook receiver ------------------------------------------
    #
    # OpenLoom's outbound webhook posts a JSON event for every task lifecycle
    # change. We only care about terminal events (TASK_COMPLETED / TASK_FAILED)
    # because they trigger result write-back. Intermediate TASK_UPDATED events
    # are accepted (logged) but do not produce a result file.
    #
    # The receiver runs in a thread (stdlib http.server is blocking) inside an
    # asyncio task via ``asyncio.to_thread``. The handler invokes
    # ``write_result`` synchronously -- which is fine because all of its work
    # is short-lived Connector I/O and it does not touch the asyncio loop.

    async def _receiver_loop(self) -> None:
        server = self._build_receiver_server()
        self._receiver_server = server
        try:
            await asyncio.to_thread(server.serve_forever)
        except Exception:
            _logger.exception("receiver crashed")
            raise
        finally:
            self._receiver_server = None
            server.server_close()

    def _build_receiver_server(self) -> ThreadingHTTPServer:
        cfg = self._config.outbound
        runner_ref = self

        class _Handler(BaseHTTPRequestHandler):
            # Quieter logs -- BaseHTTPRequestHandler logs every request to
            # stderr by default, which is noisy when OpenLoom polls every 8s.
            def log_message(self, format: str, *args: Any) -> None:  # noqa: A002
                _logger.debug("receiver " + format, *args)

            def do_POST(self) -> None:  # noqa: N802 (BaseHTTPRequestHandler API)
                if self.path != cfg.path:
                    self.send_error(404, "unknown path")
                    return
                length = int(self.headers.get("Content-Length") or 0)
                try:
                    raw = self.rfile.read(length) if length else b""
                    event = json.loads(raw.decode("utf-8"))
                except (ValueError, UnicodeDecodeError) as exc:
                    _logger.warning("receiver: invalid body: %s", exc)
                    self.send_error(400, "invalid json")
                    return

                event_name = str(event.get("event") or "")
                task_id = str(event.get("task_id") or "")
                if not event_name or not task_id:
                    _logger.warning(
                        "receiver: missing event/task_id: %r", event,
                    )
                    self.send_error(400, "missing event or task_id")
                    return

                # Only terminal events trigger write_result. Everything else is
                # acknowledged but ignored -- we just don't have anything to
                # write back yet.
                if event_name not in ("TASK_COMPLETED", "TASK_FAILED"):
                    _logger.debug(
                        "receiver: ignoring event %s for %s", event_name, task_id,
                    )
                    self._ok({"ok": True, "ignored": event_name})
                    return

                if task_id not in runner_ref._task_to_file:
                    _logger.info(
                        "receiver: %s for unknown task %s -- nothing to write",
                        event_name, task_id,
                    )
                    self._ok({"ok": True, "ignored": "unknown task"})
                    return

                data = event.get("data") or {}
                task_name = str(event.get("task_name") or "")
                # event_name is one of TASK_COMPLETED / TASK_FAILED here
                status = str(data.get("status") or event_name[len("TASK_"):].lower())
                try:
                    runner_ref.write_result(
                        task_id=task_id,
                        task_name=task_name,
                        status=status,
                        data=data if isinstance(data, dict) else {"raw": data},
                    )
                    self._ok({"ok": True, "task_id": task_id, "status": status})
                except Exception as exc:
                    _logger.exception("write_result failed for %s", task_id)
                    self.send_error(500, f"write_result failed: {exc}")

            def _ok(self, body: dict[str, Any]) -> None:
                payload = json.dumps(body).encode("utf-8")
                self.send_response(200)
                self.send_header("Content-Type", "application/json")
                self.send_header("Content-Length", str(len(payload)))
                self.end_headers()
                self.wfile.write(payload)

        # ThreadingHTTPServer handles one request per thread, so a slow
        # write_result cannot block subsequent events.
        return ThreadingHTTPServer((cfg.host, cfg.port), _Handler)

    # ── poll cycle ──────────────────────────────────────────────────────

    def _poll_once(self) -> None:
        """One polling cycle — list inbox, dispatch unseen task files."""
        entries = self._connector.list_inbox()
        for entry in entries:
            if entry.path in self._seen:
                continue
            name = PurePosixPath(entry.path).name
            if not name.startswith(self._config.task_prefix):
                continue
            ext = PurePosixPath(entry.path).suffix.lower()
            if ext not in _TASK_EXTENSIONS:
                continue
            self._dispatch(entry)

    def _dispatch(self, entry: FileEntry) -> None:
        """Download one file, push to OpenLoom, register for completion."""
        content = self._connector.download(entry.path)
        if content is None:
            return
        spec = _parse_spec(content, entry.path)
        if spec is None:
            _logger.warning("skipping %s — not a valid task spec", entry.path)
            self._seen.add(entry.path)
            return
        goal = str(spec.get("goal") or "").strip()
        if not goal:
            _logger.warning("skipping %s — missing 'goal'", entry.path)
            self._seen.add(entry.path)
            return
        workspace = str(spec.get("workspace") or spec.get("cwd") or "").strip()
        session_id = str(spec.get("sessionId") or spec.get("session_id") or "").strip()
        if not workspace and not session_id:
            _logger.warning(
                "skipping %s — need workspace or sessionId", entry.path,
            )
            self._seen.add(entry.path)
            return

        task_id = self._push_to_openloom(spec, entry)
        if task_id:
            self._seen.add(entry.path)
            self._task_to_file[task_id] = entry.path
            _logger.info(
                "pushed task %s from %s (workspace=%r sessionId=%r)",
                task_id, entry.path, workspace, session_id,
            )

    def _push_to_openloom(self, spec: dict[str, Any], entry: FileEntry) -> str | None:
        """POST task spec to OpenLoom webhook. Return task_id or None."""
        url = self._config.webhook.url
        if not url:
            _logger.error("webhook url not configured")
            return None
        payload = json.dumps(spec).encode()
        headers = {"Content-Type": "application/json"}
        if self._config.webhook.signing_secret:
            sig = _sign_payload(self._config.webhook.signing_secret, payload)
            headers["X-OpenLoom-Signature-256"] = f"sha256={sig}"

        # ``trust_env=False`` makes httpx ignore HTTP_PROXY / HTTPS_PROXY /
        # ALL_PROXY and the OS proxy config. Without it, a system-wide proxy
        # (corporate VPN, Clash, mitmproxy, etc.) hijacks requests to
        # 127.0.0.1 and the connector ends up POSTing through someone else's
        # gateway — which returns a generic "Content Filter - Access Denied"
        # page instead of the real OpenLoom response.
        try:
            with httpx.Client(timeout=10.0, trust_env=False) as client:
                resp = client.post(url, content=payload, headers=headers)
        except httpx.HTTPError as exc:
            _logger.error("POST %s failed: %s", url, exc)
            return None

        if resp.status_code >= 400:
            _logger.error(
                "POST %s returned %s: %s", url, resp.status_code, resp.text[:200],
            )
            return None
        try:
            data = resp.json()
        except Exception:
            return None
        return data.get("taskId") or data.get("task_id")

    # ── result write-back ───────────────────────────────────────────────

    def write_result(
        self,
        task_id: str,
        task_name: str,
        status: str,
        data: dict[str, Any],
    ) -> None:
        """Write a task result back to the connector's outbox.

        The result is written in the same format as the input file
        (``task-X.json`` → ``task-X.result.json`` etc.).
        """
        source_file = self._task_to_file.pop(task_id, None)
        if source_file is None:
            return

        source_ext = PurePosixPath(source_file).suffix.lower()
        result_suffix = _RESULT_SUFFIX_BY_EXT.get(source_ext, ".result.json")
        source_stem = _strip_prefix(
            PurePosixPath(source_file).name,
            self._config.task_prefix,
        )
        source_stem = PurePosixPath(source_stem).stem
        out_name = f"{self._config.task_prefix}{source_stem}{result_suffix}"
        out_path = f"{self._config.outbox_dir}/{out_name}"

        result = {
            "schema_version": "1.0",
            "task_id": task_id,
            "task_name": task_name,
            "status": status,
            "timestamp": time.time(),
            "data": data,
        }

        try:
            content = _render_result(result, source_ext)
            self._connector.upload(out_path, content)
            _logger.info("wrote result to %s", out_path)
        except Exception:
            _logger.exception("upload failed for %s", out_path)
            return

        # Archive + delete the consumed input file
        if self._config.archive_dir:
            archive_path = f"{self._config.archive_dir}/{PurePosixPath(source_file).name}"
            try:
                original = self._connector.download(source_file)
                if original is not None:
                    self._connector.upload(archive_path, original)
            except Exception:
                _logger.exception("archive failed for %s", archive_path)

        try:
            self._connector.delete_inbox(source_file)
        except Exception:
            _logger.exception("delete_inbox failed for %s", source_file)


# ── spec parsing ──────────────────────────────────────────────────────────


def _parse_spec(raw: bytes, filepath: str) -> dict[str, Any] | None:
    """Parse a task file (json/yaml/docx) into a normalized spec dict."""
    ext = PurePosixPath(filepath).suffix.lower()
    try:
        if ext == ".json":
            data = json.loads(raw)
            return data if isinstance(data, dict) else None
        if ext in (".yaml", ".yml"):
            import yaml
            data = yaml.safe_load(raw)
            return data if isinstance(data, dict) else None
        if ext == ".docx":
            return _parse_docx(raw)
    except Exception:
        return None
    return None


def _parse_docx(raw: bytes) -> dict[str, Any] | None:
    """Parse a docx task file: first table with 2 columns (field | value).

    Recognized field names (case-insensitive): ``goal``, ``workspace``,
    ``sessionId``, ``session_id``, ``name``. Any other field goes into
    ``metadata``.
    """
    from docx import Document  # python-docx

    doc = Document(io.BytesIO(raw))
    tables = doc.tables
    if not tables:
        return None

    spec: dict[str, Any] = {}
    metadata: dict[str, Any] = {}
    for row in tables[0].rows:
        if len(row.cells) < 2:
            continue
        key = row.cells[0].text.strip().lower()
        value = row.cells[1].text.strip()
        if not key:
            continue
        if key in ("goal", "workspace", "cwd", "name", "title", "sessionid", "session_id"):
            if key == "sessionid":
                key = "sessionId"
            spec[key] = value
        else:
            metadata[key] = value

    if metadata:
        spec["metadata"] = metadata
    return spec or None


# ── result rendering ──────────────────────────────────────────────────────


def _render_result(result: dict[str, Any], source_ext: str) -> bytes:
    """Render a result in the same format as the input."""
    if source_ext == ".json":
        return json.dumps(result, indent=2, ensure_ascii=False).encode()
    if source_ext in (".yaml", ".yml"):
        import yaml
        return yaml.safe_dump(result, allow_unicode=True).encode()
    if source_ext == ".docx":
        return _render_docx_result(result).read()
    return json.dumps(result, indent=2).encode()


# ── docx rendering ────────────────────────────────────────────────────────


_STATUS_HEADING = {
    "completed": "任务完成报告",
    "failed": "任务失败报告",
    "running": "任务进行中",
    "waiting": "任务等待中",
    "archived": "任务已归档",
}


def _render_docx_result(result: dict[str, Any]) -> io.BytesIO:
    """Render a result as a structured docx for phone-friendly reading.

    Layout (four sections):

      1. Header — task_name + status (heading) + timestamp
      2. Summary — the agent's ``summary`` text (one paragraph)
      3. Agent trace — each recent_activity entry becomes its own
         block with timestamp + text + tool calls (each tool on its
         own line, indented). Tools are visually separated from the
         prose so they do not mix into the agent's narrative.
      4. Metadata — anything left over (e.g. active_session_id) in a
         small two-column table.

    All content is preserved verbatim — no truncation. Long lines
    wrap naturally in Word / WPS on the phone.
    """
    from docx import Document

    doc = Document()

    status = str(result.get("status") or "").lower()
    heading_text = _STATUS_HEADING.get(status, "OpenLoom Task Result")
    doc.add_heading(heading_text, level=1)

    task_name = str(result.get("task_name") or result.get("task_id") or "")
    if task_name:
        doc.add_heading(task_name, level=2)

    _docx_add_metadata_block(doc, result)

    data = result.get("data") or {}
    if isinstance(data, dict):
        summary = str(data.get("summary") or "").strip()
        if summary:
            doc.add_heading("概要", level=2)
            doc.add_paragraph(summary)

        recent = data.get("recent_activity")
        if isinstance(recent, list) and recent:
            _docx_add_trace(doc, recent)

        leftover = _docx_leftover(data)
        if leftover:
            doc.add_heading("其他信息", level=2)
            table = doc.add_table(rows=0, cols=2)
            for k, v in leftover.items():
                row = table.add_row().cells
                row[0].text = str(k)
                row[1].text = str(v)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def _docx_add_metadata_block(doc: Any, result: dict[str, Any]) -> None:
    """Top-level status banner: task_id, status, timestamp."""
    from datetime import UTC, datetime

    lines: list[str] = []
    status = str(result.get("status") or "")
    task_id = str(result.get("task_id") or "")
    timestamp = result.get("timestamp")
    if status:
        lines.append(f"状态: {status}")
    if task_id:
        lines.append(f"任务 ID: {task_id}")
    if isinstance(timestamp, (int, float)) and timestamp > 0:
        iso = datetime.fromtimestamp(timestamp, tz=UTC).strftime("%Y-%m-%d %H:%M:%S")
        lines.append(f"时间: {iso}")

    if lines:
        para = doc.add_paragraph()
        for line in lines:
            run = para.add_run(line + "\n")
            run.font.size = None
        para.runs[0].bold = True


_TOOL_STATUS_GLYPH = {
    "completed": "✓",
    "running": "▶",
    "pending": "…",
    "active": "▶",
    "failed": "✗",
    "error": "✗",
}


def _docx_add_trace(doc: Any, recent: list[Any]) -> None:
    """Render recent_activity as a separate block per entry.

    Each entry is its own paragraph: a one-line heading (counter +
    timestamp + status), then the agent text, then a list of tool
    calls. Tools are rendered as indented bullet items so they read
    as a separate list from the agent's prose.
    """
    from datetime import UTC, datetime

    doc.add_heading(f"Agent 执行轨迹（共 {len(recent)} 条）", level=2)

    for idx, entry in enumerate(recent, start=1):
        if not isinstance(entry, dict):
            continue
        ts = ""
        completed = entry.get("completed_at")
        if isinstance(completed, (int, float)) and completed > 0:
            ts = datetime.fromtimestamp(
                completed / 1000 if completed > 1e12 else completed,
                tz=UTC,
            ).strftime("%H:%M:%S")

        head = doc.add_paragraph()
        head_run = head.add_run(f"{idx}. {ts}".rstrip())
        head_run.bold = True

        text = str(entry.get("text") or "").strip()
        if text:
            doc.add_paragraph(text)

        tools = entry.get("tools") or []
        if isinstance(tools, list) and tools:
            for tool in tools:
                if not isinstance(tool, dict):
                    continue
                _docx_add_tool_line(doc, tool)


def _docx_add_tool_line(doc: Any, tool: dict[str, Any]) -> None:
    """One bullet for a single tool call, glyph by status."""
    name = str(tool.get("tool") or tool.get("name") or "tool")
    status = str(tool.get("status") or "unknown").lower()
    glyph = _TOOL_STATUS_GLYPH.get(status, "•")
    excerpt = str(tool.get("input_excerpt") or "")
    line = f"    {glyph} {name} [{status}]"
    if excerpt:
        line += f"  {excerpt}"

    para = doc.add_paragraph()
    para.paragraph_format.left_indent = None
    run = para.add_run(line)
    run.font.name = "Consolas"


def _docx_leftover(data: dict[str, Any]) -> dict[str, Any]:
    """Return data keys that were not already promoted to a section."""
    used = {"summary", "recent_activity"}
    return {k: v for k, v in data.items() if k not in used}


# ── legacy helpers (kept for callers that still depend on the flatten table) ─


def _flatten_for_docx(result: dict[str, Any]) -> list[tuple[str, str]]:
    """Deprecated. Use ``_render_docx_result`` directly.

    Kept for tests that still assert the old 2-column row shape; new
    consumers should not depend on this.
    """
    flat: list[tuple[str, str]] = []
    for key in ("schema_version", "task_id", "task_name", "status"):
        if key in result:
            flat.append((key, str(result[key])))
    if "timestamp" in result:
        from datetime import UTC, datetime
        ts = result["timestamp"]
        flat.append(("timestamp", str(ts)))
        flat.append(("timestamp_iso", datetime.fromtimestamp(ts, tz=UTC).strftime("%Y-%m-%dT%H:%M:%SZ")))
    data = result.get("data") or {}
    if isinstance(data, dict):
        for k, v in data.items():
            flat.append((f"data.{k}", str(v)))
    elif data:
        flat.append(("data", str(data)))
    return flat


# ── helpers ───────────────────────────────────────────────────────────────


def _sign_payload(secret: str, payload: bytes) -> str:
    return hmac.new(
        secret.encode("utf-8"), payload, hashlib.sha256,
    ).hexdigest()


def _strip_prefix(name: str, prefix: str) -> str:
    return name[len(prefix):] if name.startswith(prefix) else name
